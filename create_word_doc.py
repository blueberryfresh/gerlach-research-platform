"""
Create a properly formatted Word document from the Participant Instructions
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_formatted_word_doc():
    """Create a well-formatted Word document"""
    
    # Read the markdown content
    md_file = Path('docs/PARTICIPANT_INSTRUCTIONS.md')
    content = md_file.read_text(encoding='utf-8')
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add spacing
            doc.add_paragraph()
        
        elif line.startswith('# '):
            # Main heading (H1)
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line.startswith('## '):
            # Section heading (H2)
            doc.add_heading(line[3:], level=2)
        
        elif line.startswith('### '):
            # Subsection heading (H3)
            doc.add_heading(line[4:], level=3)
        
        elif line.startswith('#### '):
            # Sub-subsection heading (H4)
            doc.add_heading(line[5:], level=4)
        
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        
        elif line.startswith('✅ ') or line.startswith('❌ ') or line.startswith('⚠️ '):
            # Emoji bullet points
            doc.add_paragraph(line, style='List Bullet')
        
        elif line.startswith('1. ') or (len(line) > 2 and line[0].isdigit() and line[1:3] == '. '):
            # Numbered list
            doc.add_paragraph(line[3:], style='List Number')
        
        elif line.startswith('```'):
            # Code block marker - skip
            continue
        
        elif line.startswith('|'):
            # Table row - handle specially
            # For now, just add as paragraph
            doc.add_paragraph(line, style='Normal')
        
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        
        elif line.startswith('---'):
            # Horizontal rule - add spacing
            doc.add_paragraph()
            p = doc.add_paragraph('_' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save document
    output_file = 'PARTICIPANT_INSTRUCTIONS.docx'
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    print(f"📍 Location: {Path(output_file).absolute()}")
    return output_file

if __name__ == '__main__':
    create_formatted_word_doc()
